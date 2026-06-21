from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "report_assets"
OUTPUT = ROOT / "docs" / "Weather_Decision_Agent_Final_Report.docx"
DEPLOYMENT_SCREENSHOT = Path(
    "/var/folders/pn/0jbn8xs174524_3ngf9vs_2m0000gn/T/TemporaryItems/"
    "NSIRD_screencaptureui_9JcqEL/Ekran Resmi 2026-06-21 13.35.42.png"
)

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x27, 0x30)
MUTED = RGBColor(0x62, 0x6B, 0x75)
ACCENT = RGBColor(0xE6, 0x62, 0x57)
LIGHT_FILL = "F2F4F7"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_element = table._tbl
    table_pr = table_element.tblPr

    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        body = paragraph.add_run(text[len(bold_lead):])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_figure(doc, path, caption, width=6.25):
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(4)
    caption_paragraph.paragraph_format.space_after = Pt(10)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED, italic=True)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Weather Decision Agent  |  Final Project Report")
    set_run_font(header_run, size=8.5, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def build_cover(doc):
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run("FINAL PROJECT REPORT")
    set_run_font(kicker_run, size=10, color=ACCENT, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Weather Decision Agent")
    set_run_font(title_run, size=30, color=DARK_BLUE, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "A weather-aware activity and venue recommendation system"
    )
    set_run_font(subtitle_run, size=14, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(46)

    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team.add_run("Deniz Özmen (2203032)  |  Ömer Şahin (2104101)")
    set_run_font(team_run, size=11.5, color=INK, bold=True)
    team.paragraph_format.space_after = Pt(8)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run("June 2026")
    set_run_font(date_run, size=11, color=MUTED)
    date.paragraph_format.space_after = Pt(34)

    demo = doc.add_paragraph()
    demo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(
        demo,
        "Live Demo",
        "https://dozmen23-weather-decision-agent.hf.space/",
    )
    demo.add_run("   |   ")
    add_hyperlink(
        demo,
        "GitHub Repository",
        "https://github.com/dozmen23/weather-decision-agent-v2",
    )
    doc.add_page_break()


def add_initial_report(doc):
    add_heading(doc, "1. Motivation", 1)
    add_body(
        doc,
        "Modern weather applications usually show raw values, but they still leave the final decision to the user. The Weather Decision Agent was proposed to turn temperature, rain and wind data into a simple answer: what can I comfortably do today?",
    )

    add_heading(doc, "2. Problem Definition", 1)
    add_body(
        doc,
        "This is a decision-making problem under uncertainty. Weather conditions change, several variables must be considered together, and two users may react differently to the same forecast. The project therefore combines decision support, context-aware recommendation and agentic AI concepts.",
    )

    add_heading(doc, "3. Project Overview", 1)
    add_body(
        doc,
        "The original plan was a Python and Streamlit application that would collect weather data, read user preferences, evaluate a catalog of activities and return ranked recommendations with short explanations.",
    )
    add_body(
        doc,
        "The first technical design expected OpenWeatherMap and a GPT-4o mini based decision agent. During development, this structure was deliberately changed to make the system safer, easier to test and less dependent on the language model.",
    )

    add_heading(doc, "4. Initial System Workflow", 1)
    add_body(
        doc,
        "The planned workflow started with location and preference inputs. Weather data would then be normalized, combined with activity options and sent to the decision agent. The final output would contain ranked activities and a short reason for each recommendation.",
    )

    add_heading(doc, "5. Team Responsibilities", 1)
    add_body(
        doc,
        "Deniz Özmen focused on the decision workflow, recommendation rules, LLM integration and the preference experience. Ömer Şahin focused on weather data collection, normalization, structured activity data and connecting the service layer to the Streamlit output.",
    )


def add_comparison_table(doc):
    add_heading(doc, "6. From Initial Plan to Final Product", 1)
    add_body(
        doc,
        "The project kept its original goal, but the implementation became more controlled and much closer to a real product.",
    )
    rows = [
        ("Area", "Initial plan", "Final implementation"),
        ("Weather", "OpenWeatherMap", "Open-Meteo, 7-day forecast"),
        ("Main decision", "LLM-based ranking", "Deterministic rules and scoring"),
        ("LLM role", "Choose activities", "Explain and review only"),
        ("Location", "City input", "City or Google Maps selection"),
        ("Venues", "Static activity options", "Verified Google Places results"),
        ("Deployment", "Local Streamlit app", "Docker Space on Hugging Face"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    row_properties.append(header_marker)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5, bold=row_index == 0)
    set_table_geometry(table, [1800, 3300, 4260])


def add_final_implementation(doc):
    doc.add_page_break()
    add_heading(doc, "7. Final Implementation Update", 1)
    add_body(
        doc,
        "The final version is a hybrid decision system. Weather safety, scoring and fallback rules are deterministic. The LLM cannot change a safety rule or score; it only turns the result into a natural explanation and acts as a second reviewer.",
    )

    add_heading(doc, "7.1 Weather Risk and Scoring", 2)
    add_body(
        doc,
        "Rain, wind, temperature and weather condition are combined into four user-friendly risk labels: comfortable, cautious, risky and very risky. Recommendations also keep a clear score breakdown for weather safety, preference match, comfort and practicality.",
    )

    add_heading(doc, "7.2 Smarter Fallback Logic", 2)
    add_body(
        doc,
        "The system first searches for an exact activity match. If the weather makes that option unsafe, it looks for a close indoor alternative instead of returning a random activity. For example, an outdoor walk can become an indoor track or mall walk.",
    )

    add_heading(doc, "8. User Experience", 1)
    add_body(
        doc,
        "User Mode keeps the screen simple: seven forecast cards, understandable preference controls and a short explanation written in everyday language. Developer Mode hides behind a separate option and shows the trace, scoring details, raw weather data and evaluator output when technical inspection is needed.",
    )
    add_figure(
        doc,
        ASSETS / "final_home.png",
        "Figure 1. Final User Mode with seven-day forecast cards and simple preferences.",
    )

    doc.add_page_break()
    add_heading(doc, "9. Map and Real Venue Integration", 1)
    add_body(
        doc,
        "Users can enter a city or select any point directly on Google Maps. The selected coordinates are used both for the weather forecast and for nearby venue searches. Real venue results come from Google Places; the LLM is not allowed to invent a location.",
    )
    add_figure(
        doc,
        ASSETS / "final_map.png",
        "Figure 2. Google Maps location picker running in the deployed application.",
    )

    doc.add_page_break()
    add_heading(doc, "10. Recommendation Output", 1)
    add_body(
        doc,
        "The result explains the weather first, then shows the selected activity, why it fits and what the user should pay attention to. When available, verified nearby venues are displayed with distance, accessibility and a direct Google Maps link.",
    )
    add_figure(
        doc,
        ASSETS / "final_recommendation.png",
        "Figure 3. Final recommendation with friendly explanation and verified venue results.",
    )

    add_heading(doc, "11. Evaluation and Safety", 1)
    add_body(
        doc,
        "The project includes automated tests and reusable evaluation scenarios for thunderstorms, high wind, extreme temperature, light rain, exact preference matches and cases where no safe activity exists. The current suite contains 159 passing tests.",
    )
    add_body(
        doc,
        "LLM safety checks also cover invented activities, changed scores and unsafe suggestions. If the model output is invalid, the deterministic recommendation remains the source of truth.",
    )

    doc.add_page_break()
    add_heading(doc, "12. Deployment", 1)
    add_body(
        doc,
        "The application is packaged with Docker and deployed publicly on Hugging Face Spaces. GitHub Actions automatically synchronizes every push to the main branch, so approved code changes can reach the live demo without a manual upload.",
    )
    add_figure(
        doc,
        DEPLOYMENT_SCREENSHOT,
        "Figure 4. Hugging Face container logs showing the Streamlit service on port 7860.",
    )

    add_heading(doc, "13. Final Architecture", 1)
    add_body(
        doc,
        "In simple terms, the user chooses a place and preferences, Open-Meteo provides the forecast, the deterministic agent selects safe activities, Google Places returns real venues and the LLM explains the already-approved result. Streamlit brings these parts together in one interface.",
    )

    add_heading(doc, "14. Conclusion", 1)
    add_body(
        doc,
        "The final Weather Decision Agent goes beyond displaying weather. It turns a forecast into a practical plan while keeping safety decisions transparent and testable. The result is a small but complete agentic AI product: useful for the user, inspectable for the developer and available through a public demo link.",
    )

    links = doc.add_paragraph()
    links.paragraph_format.space_before = Pt(12)
    lead = links.add_run("Project links: ")
    set_run_font(lead, bold=True)
    add_hyperlink(
        links,
        "Live Demo",
        "https://dozmen23-weather-decision-agent.hf.space/",
    )
    links.add_run("  |  ")
    add_hyperlink(
        links,
        "GitHub",
        "https://github.com/dozmen23/weather-decision-agent-v2",
    )


def main():
    document = Document()
    configure_document(document)
    build_cover(document)
    add_initial_report(document)
    add_comparison_table(document)
    add_final_implementation(document)
    document.core_properties.title = "Weather Decision Agent - Final Project Report"
    document.core_properties.subject = "Final implementation and deployment report"
    document.core_properties.author = "Deniz Özmen and Ömer Şahin"
    document.core_properties.keywords = "Weather, Agentic AI, Streamlit, Google Maps, Open-Meteo"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

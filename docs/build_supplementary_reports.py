from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_final_report import (
    ACCENT,
    BLUE,
    DARK_BLUE,
    INK,
    LIGHT_FILL,
    MUTED,
    add_body,
    add_figure,
    add_heading,
    add_hyperlink,
    add_page_field,
    set_cell_margins,
    set_cell_shading,
    set_run_font,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
EXECUTIVE_OUTPUT = DOCS / "Weather_Decision_Agent_Executive_Summary.docx"
PRACTITIONER_OUTPUT = DOCS / "Weather_Decision_Agent_Practitioner_Notes.docx"


def configure_document(doc, preset):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    if preset == "standard_business_brief":
        body_after = 6
        line_spacing = 1.10
        heading_tokens = {
            "Heading 1": (16, BLUE, 16, 8),
            "Heading 2": (13, BLUE, 12, 6),
            "Heading 3": (12, DARK_BLUE, 8, 4),
        }
    else:
        body_after = 6
        line_spacing = 1.25
        heading_tokens = {
            "Heading 1": (16, BLUE, 18, 10),
            "Heading 2": (13, BLUE, 14, 7),
            "Heading 3": (12, DARK_BLUE, 10, 5),
        }

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(body_after)
    normal.paragraph_format.line_spacing = line_spacing

    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    add_page_field(section.footer.paragraphs[0])


def set_header(doc, left, right):
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = header.add_run(left)
    set_run_font(left_run, size=8.5, color=MUTED, bold=True)
    right_run = header.add_run(f"  |  {right}")
    set_run_font(right_run, size=8.5, color=MUTED)


def add_metadata_line(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.5)


def add_lead(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    paragraph_properties.append(shading)
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=11, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11)


def mark_header_row(row):
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def add_table(doc, rows, widths_dxa, font_size=9.5):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    mark_header_row(table.rows[0])
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=font_size, bold=row_index == 0)
    set_table_geometry(table, widths_dxa)
    return table


def add_code_block(doc, lines):
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0 if index < len(lines) - 1 else 8)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F4F7")
        properties.append(shading)
        run = paragraph.add_run(line)
        set_run_font(run, name="Courier New", size=9.5, color=INK)


def add_title_block(doc, kicker, title, subtitle):
    kicker_paragraph = doc.add_paragraph()
    kicker_run = kicker_paragraph.add_run(kicker.upper())
    set_run_font(kicker_run, size=9.5, color=ACCENT, bold=True)
    kicker_paragraph.paragraph_format.space_after = Pt(5)

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, size=25, color=DARK_BLUE, bold=True)
    title_paragraph.paragraph_format.space_after = Pt(5)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_run_font(subtitle_run, size=12.5, color=MUTED)
    subtitle_paragraph.paragraph_format.space_after = Pt(18)


def build_executive_summary():
    doc = Document()
    configure_document(doc, "standard_business_brief")
    set_header(doc, "Executive Summary", "Weather Decision Agent")
    add_title_block(
        doc,
        "Executive Summary",
        "Weather Decision Agent",
        "A weather-aware activity and venue recommendation system",
    )
    add_metadata_line(doc, "Team", "Deniz Özmen (2203032), Ömer Şahin (2104101)")
    add_metadata_line(doc, "Status", "Implemented, tested and publicly deployed")
    add_metadata_line(doc, "Date", "June 2026")

    add_lead(
        doc,
        "Executive statement.",
        "The project converts a seven-day weather forecast and personal preferences into safe, understandable activity plans. It is available as a public web application and supports real location selection and verified venue results.",
    )

    add_heading(doc, "The Challenge", 1)
    add_body(
        doc,
        "Most weather applications provide data but leave the interpretation to the user. Temperature, rain and wind must be considered together, and the same conditions may be acceptable for one activity but unsafe for another. The project addresses this gap as a decision-support problem rather than a weather-display problem.",
    )

    add_heading(doc, "The Delivered Solution", 1)
    add_body(
        doc,
        "Users select a city or a point on Google Maps, choose a day and describe their activity preferences. Open-Meteo supplies the forecast. A deterministic decision layer evaluates weather safety, preference fit, comfort and practicality before producing recommendations. Google Places adds verified nearby venues when relevant.",
    )
    add_body(
        doc,
        "The interface separates a simple User Mode from a technical Developer Mode. User Mode explains why an activity fits and what to watch for. Developer Mode exposes the score breakdown, rule trace, evaluator results and raw weather data.",
    )

    add_heading(doc, "Why the Architecture Matters", 1)
    add_body(
        doc,
        "The initial plan placed the language model at the center of activity selection. The final architecture intentionally moved safety and scoring into deterministic code. The LLM is limited to explanation, controlled candidate generation and second-review support. It cannot change safety thresholds, scores or verified venue data.",
    )

    add_heading(doc, "Evidence of Completion", 1)
    add_table(
        doc,
        [
            ("Area", "Result"),
            ("Product", "Public Streamlit application on Hugging Face Spaces"),
            ("Decision quality", "Deterministic safety, fallback and score breakdown"),
            ("Location", "City search and Google Maps coordinate selection"),
            ("Venues", "Google Places provider with verified Maps links"),
            ("Quality", "159 automated tests plus reusable evaluation scenarios"),
            ("Delivery", "Docker deployment and GitHub Actions synchronization"),
        ],
        [2200, 7160],
    )

    doc.add_page_break()
    add_heading(doc, "User-Facing Result", 1)
    add_figure(
        doc,
        ASSETS / "final_recommendation.png",
        "Final recommendation output with weather risk, user-friendly reasoning and a safe activity result.",
        width=6.1,
    )

    add_heading(doc, "Current Limits and Next Step", 1)
    add_body(
        doc,
        "Google Places category quality can vary by location, free hosting may introduce a cold-start delay and personalization currently uses a small feedback-based score adjustment rather than advanced machine learning. The most valuable next step is to observe real usage and improve activity-to-venue matching from evidence rather than expanding scope blindly.",
    )

    add_heading(doc, "Bottom Line", 1)
    add_body(
        doc,
        "Weather Decision Agent is a complete small-scale agentic AI product. It is useful to an end user, inspectable by a developer and designed so that the language model supports the decision without controlling safety-critical behavior.",
    )

    link_paragraph = doc.add_paragraph()
    link_paragraph.paragraph_format.space_before = Pt(8)
    lead = link_paragraph.add_run("Live demo: ")
    set_run_font(lead, bold=True)
    add_hyperlink(
        link_paragraph,
        "Open live demo",
        "https://dozmen23-weather-decision-agent.hf.space/",
    )

    doc.core_properties.title = "Weather Decision Agent - Executive Summary"
    doc.core_properties.author = "Deniz Özmen and Ömer Şahin"
    doc.save(EXECUTIVE_OUTPUT)


def build_practitioner_notes():
    doc = Document()
    configure_document(doc, "compact_reference_guide")
    set_header(doc, "Practitioner Notes", "Weather Decision Agent")

    for _ in range(4):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Weather Decision Agent")
    set_run_font(title_run, size=29, color=DARK_BLUE, bold=True)
    title.paragraph_format.space_after = Pt(6)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Practitioner Notes")
    set_run_font(subtitle_run, size=16, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(14)
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description_run = description.add_run(
        "A concise guide for running, configuring, testing and maintaining the project"
    )
    set_run_font(description_run, size=11, color=MUTED, italic=True)
    description.paragraph_format.space_after = Pt(50)
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team.add_run("Deniz Özmen · Ömer Şahin · June 2026")
    set_run_font(team_run, size=10.5, color=INK, bold=True)
    doc.add_page_break()

    add_heading(doc, "1. Purpose and Operating Model", 1)
    add_body(
        doc,
        "These notes are for a developer, evaluator or maintainer who needs to run the application and understand where decisions are made. The core rule is simple: deterministic code owns safety and scoring; external services provide data; the LLM explains an approved result.",
    )

    add_heading(doc, "2. System Flow", 1)
    add_table(
        doc,
        [
            ("Stage", "Responsibility", "Main component"),
            ("1", "Collect location, date and preferences", "Streamlit UI"),
            ("2", "Fetch and normalize seven-day forecast", "Open-Meteo service"),
            ("3", "Calculate risk and reject unsafe options", "Rules and scoring"),
            ("4", "Try exact match, then close fallback", "Decision agent"),
            ("5", "Find verified nearby venues", "Venue provider / Google Places"),
            ("6", "Explain the approved recommendation", "Optional LLM services"),
        ],
        [900, 4940, 3520],
        font_size=9,
    )

    add_heading(doc, "3. Local Setup", 1)
    add_body(doc, "Use Python 3.12 and install the pinned project dependencies.")
    add_code_block(
        doc,
        [
            "python3.12 -m venv .venv",
            "source .venv/bin/activate",
            "pip install -r requirements.txt",
        ],
    )
    add_body(
        doc,
        "Copy `.env.example` to `.env` and fill only the services that will be enabled. Never commit `.env` or real API keys.",
    )

    add_heading(doc, "4. Environment Configuration", 1)
    add_table(
        doc,
        [
            ("Variable", "When needed", "Purpose"),
            ("VENUE_PROVIDER", "Always", "Select `json` or `google_places`"),
            ("GOOGLE_PLACES_API_KEY", "Google Places mode", "Server-side venue search secret"),
            ("GOOGLE_MAPS_BROWSER_API_KEY", "Google Maps picker", "Referrer-restricted browser key"),
            ("LLM_ENABLED", "Always", "Enable or disable LLM assistance"),
            ("LLM_PROVIDER", "LLM enabled", "Select the configured LLM provider"),
            ("LLM_MODEL", "LLM enabled", "Choose the configured model"),
            ("LLM_API_KEY", "LLM enabled", "Private provider credential"),
        ],
        [2650, 2250, 4460],
        font_size=8.8,
    )

    add_heading(doc, "5. Running the Application", 1)
    add_code_block(doc, ["streamlit run streamlit_app.py"])
    add_body(
        doc,
        "Start in User Mode for the normal recommendation flow. Use Developer Mode only when inspecting rule traces, score components, evaluator results, raw weather data or provider behavior.",
    )

    add_heading(doc, "6. Decision Behavior", 1)
    add_heading(doc, "6.1 Risk", 2)
    add_body(
        doc,
        "Rain, wind, temperature and weather condition are converted into LOW, MODERATE, HIGH or SEVERE risk. User Mode translates these into comfortable, cautious, risky and very risky labels.",
    )
    add_heading(doc, "6.2 Matching and Fallback", 2)
    add_body(
        doc,
        "The agent tries an exact activity match first. If that option is unsafe or unavailable, it searches for a close alternative. Outdoor walking should therefore fall back to an indoor walking option before an unrelated activity.",
    )
    add_heading(doc, "6.3 Scoring", 2)
    add_body(
        doc,
        "The total score is built from weather safety, preference match, comfort and practicality. A recommendation must pass deterministic safety checks before its score or explanation is shown.",
    )

    add_heading(doc, "7. Maps and Venue Providers", 1)
    add_body(
        doc,
        "The venue layer uses a provider interface. JSON data supports deterministic demos and tests; Google Places supports live venues. Keep the Google Places key on the server. Restrict the browser key to Maps JavaScript API and approved HTTP referrers such as localhost and the Hugging Face Space domain.",
    )
    add_figure(
        doc,
        ASSETS / "final_map.png",
        "Google Maps picker used to provide coordinates to both weather and venue services.",
        width=6.1,
    )

    add_heading(doc, "8. LLM Boundary", 1)
    add_body(
        doc,
        "The LLM may summarize weather, explain a recommendation, generate controlled catalog candidates and act as a second reviewer. Every candidate is revalidated. The LLM must not change risk levels, safety rules, score components or verified venue identity.",
    )

    add_heading(doc, "9. Testing and Evaluation", 1)
    add_code_block(doc, ["pytest"])
    add_body(
        doc,
        "The current suite contains 159 passing tests. Evaluation scenarios cover thunderstorms, high wind, extreme heat, light rain, exact preference matches, safe stopping, coordinate-based venue sorting and rejection of unsafe or invented LLM output.",
    )

    add_heading(doc, "10. Deployment Workflow", 1)
    add_body(
        doc,
        "The repository contains a Docker configuration for port 7860. The public application runs on Hugging Face Spaces. GitHub Actions synchronizes pushes from `main` to the Space repository.",
    )
    add_body(
        doc,
        "Store Google and LLM credentials as Hugging Face Secrets. Store non-sensitive switches such as `VENUE_PROVIDER`, `LLM_ENABLED`, `LLM_PROVIDER` and `LLM_MODEL` as Variables. A secret or variable change restarts the Space.",
    )

    add_heading(doc, "11. Operational Notes", 1)
    add_table(
        doc,
        [
            ("Situation", "Expected handling"),
            ("Weather API unavailable", "Show a controlled error; do not fabricate weather"),
            ("No safe activity", "Stop safely and explain that no option passed"),
            ("Google Places unavailable", "Keep activity result; omit live venue suggestions"),
            ("LLM unavailable or invalid", "Use deterministic explanation and result"),
            ("Free Space is sleeping", "Allow a short cold-start delay"),
        ],
        [3000, 6360],
    )

    add_heading(doc, "12. Known Limits", 1)
    add_body(
        doc,
        "Venue categories depend on Google Places coverage and may occasionally produce a weak match. Forecast selection is limited to the available seven-day window. Public demo history is session-based and temporary. Personalization is intentionally small and feedback-based rather than an advanced machine-learning model.",
    )

    add_heading(doc, "13. Maintenance Priorities", 1)
    add_table(
        doc,
        [
            ("Priority", "Practical action"),
            ("Activity quality", "Review weak activity-to-venue mappings from real examples"),
            ("Safety", "Add an evaluation case before changing a threshold"),
            ("LLM behavior", "Keep prompts short and revalidate every generated candidate"),
            ("Credentials", "Rotate keys when exposed and keep API restrictions narrow"),
            ("Deployment", "Check GitHub Actions and Space logs after a release"),
        ],
        [2600, 6760],
    )

    add_heading(doc, "14. Quick Reference", 1)
    add_code_block(
        doc,
        [
            "streamlit run streamlit_app.py   # Run locally",
            "pytest                           # Run tests",
            "git push origin main             # Trigger GitHub/HF sync",
        ],
    )

    link_paragraph = doc.add_paragraph()
    lead = link_paragraph.add_run("Live demo: ")
    set_run_font(lead, bold=True)
    add_hyperlink(
        link_paragraph,
        "Open live demo",
        "https://dozmen23-weather-decision-agent.hf.space/",
    )

    doc.core_properties.title = "Weather Decision Agent - Practitioner Notes"
    doc.core_properties.author = "Deniz Özmen and Ömer Şahin"
    doc.save(PRACTITIONER_OUTPUT)


def main():
    build_executive_summary()
    build_practitioner_notes()
    print(EXECUTIVE_OUTPUT)
    print(PRACTITIONER_OUTPUT)


if __name__ == "__main__":
    main()
